"""File -> plain text. One function per format, one dispatcher."""
from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def _read_pdf(path: Path) -> Tuple[str, List[str]]:
    from pypdf import PdfReader

    warnings: List[str] = []
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:                      # pragma: no cover
            warnings.append(f"page {i + 1} failed to extract: {exc}")
    text = "\n".join(pages)
    if not text.strip():
        # The single most common real-world failure: a scanned/image-only PDF.
        # We surface it rather than silently scoring the candidate as empty.
        warnings.append(
            "no extractable text - likely a scanned image PDF; OCR is not "
            "implemented (see docs/TRADEOFFS.md)"
        )
    return text, warnings


def _read_docx(path: Path) -> Tuple[str, List[str]]:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    # Resumes very often put skills or dates in tables, which
    # document.paragraphs does not reach.
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts), []


def _read_text(path: Path) -> Tuple[str, List[str]]:
    return path.read_text(encoding="utf-8", errors="replace"), []


_READERS = {".pdf": _read_pdf, ".docx": _read_docx,
            ".txt": _read_text, ".md": _read_text}


def extract_text(path: Path) -> Tuple[str, List[str]]:
    """Return (text, warnings). Never raises for a merely-unreadable file.

    A single corrupt resume must not abort a 200-resume batch, so failures
    are returned as warnings attached to that candidate instead.
    """
    path = Path(path)
    reader = _READERS.get(path.suffix.lower())
    if reader is None:
        return "", [f"unsupported file type '{path.suffix}'"]
    try:
        return reader(path)
    except Exception as exc:
        return "", [f"failed to parse: {type(exc).__name__}: {exc}"]


def discover_resumes(directory: Path) -> List[Path]:
    """All supported resume files in a directory, deterministically ordered."""
    directory = Path(directory)
    if not directory.is_dir():
        raise NotADirectoryError(f"resume directory not found: {directory}")
    return sorted(
        p for p in directory.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )
